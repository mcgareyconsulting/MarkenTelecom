import pandas as pd
import os
from pdf_generator.generate_pdf import ViolationNoticePDF
import re
from database.models import District, Account, ViolationReport
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Normalize address lookup
import re


def normalize_address(addr):
    if not addr:
        return ""
    addr = addr.strip().lower()
    # Remove punctuation (periods, commas, etc.)
    addr = re.sub(r"[.,]", "", addr)
    # Collapse multiple spaces
    addr = re.sub(r"\s+", " ", addr)
    # Replace common street suffixes with abbreviations
    suffix_map = {
        " street": " st",
        " avenue": " ave",
        " boulevard": " blvd",
        " drive": " dr",
        " road": " rd",
        " lane": " ln",
        " court": " ct",
        " place": " pl",
        " trail": " trl",
        " parkway": " pkwy",
        " circle": " cir",
        " terrace": " ter",
        " way": " way",
    }
    for long, short in suffix_map.items():
        if addr.endswith(long):
            addr = addr[: -len(long)] + short
        addr = addr.replace(long + " ", short + " ")
    return addr


# Collect and compare address from collected data and homeowner data
def violation_match(district_name):
    """
    Function to join district data with homeowner data based on address
    and return a structured data batch for PDF generation.

    Returns:
        list: List of dictionaries containing all data needed for PDF generation
    """
    print(f"Collecting addresses for {district_name}...")

    district = District.query.filter_by(name=district_name).first()

    if not district:
        print(f"District '{district_name}' not found in the database.")
        return []

    # Get all accounts in this district
    accounts = Account.query.filter_by(district_id=district.id).all()

    # Get all violation reports in this district
    violation_reports = ViolationReport.query.filter_by(district=district_name).all()

    # Create account lookup by normalized address
    account_lookup = {normalize_address(a.service_address): a for a in accounts}

    # Collection list for PDF generation data
    pdf_data_batch = []
    counter = 0

    # Cross-reference violation reports with accounts by address
    for report in violation_reports:
        addr_key = normalize_address(report.address_line1)
        account = account_lookup.get(addr_key)

        print(f"Checking report address: {report.address_line1} -> {addr_key}")

        if account:
            counter += 1

            # Process each violation in the report
            for violation in report.violations:

                # Collect violation images
                violation_images = []
                for image in violation.images:
                    violation_images.append(
                        {
                            "filename": image.filename,
                            "file_path": image.file_path,
                            "original_filename": image.original_filename,
                        }
                    )

                # Create data package for PDF generation
                pdf_data = {
                    # District information
                    "district_name": district.name,
                    "district_label": district.label,
                    "district_code": district.code,
                    # Homeowner information
                    "homeowner_name": account.account_name,
                    "account_number": account.account_number,
                    "lot_number": account.lot_number,
                    # Addresses
                    "property_address": report.address_line1,
                    "property_address_line2": report.address_line2,
                    "property_city": report.city,
                    "property_state": report.state,
                    "property_zip": report.zip_code,
                    "mailing_address": account.mail_address,
                    "mailing_city_st_zip": account.mail_city_st_zip,
                    "service_address": account.service_address,
                    "service_city_st_zip": account.service_city_st_zip,
                    # Contact information
                    "homeowner_email": account.email,
                    # Violation information
                    "violation_id": violation.id,
                    "violation_type": violation.violation_type,
                    "violation_notes": violation.notes,
                    "violation_date": violation.created_at.strftime("%Y-%m-%d"),
                    "violation_images": violation_images,
                    # Report information
                    "report_id": report.id,
                    "report_status": report.status,
                    "report_created_at": report.created_at.strftime("%Y-%m-%d"),
                    "report_updated_at": report.updated_at.strftime("%Y-%m-%d"),
                }
                print(violation.notes)
                pdf_data_batch.append(pdf_data)

                print(
                    f"Added to batch: {account.account_name} - {violation.violation_type}"
                )
        else:
            print(
                f"No account found for violation report address: {report.address_line1}"
            )

    print(f"Total matches found: {counter}")
    print(f"Total PDF data packages created: {len(pdf_data_batch)}")

    return pdf_data_batch

    # iterate through all violation_reports in the database with district = highlands_mead
    # and join with homeowner data based on address


def collect_violation_data_for_pdf(district_name):
    """
    Main function to collect violation data ready for PDF generation.

    Args:
        district_name (str): Name of the district to process

    Returns:
        list: List of data dictionaries ready for PDF generation
    """
    print(f"Collecting violation data for PDF generation: {district_name}")

    # Get the matched violation data
    pdf_data_batch = violation_match(district_name)

    if not pdf_data_batch:
        print("No violation data found for PDF generation.")
        return []

    # Optionally add any additional processing here
    # For example, you might want to:
    # - Filter by violation type
    # - Sort by address or date
    # - Add district-specific rules/regulations

    # Sort by property address for consistent ordering
    pdf_data_batch.sort(key=lambda x: x["property_address"])

    print(f"Prepared {len(pdf_data_batch)} violation notices for PDF generation")

    return pdf_data_batch


def convert_to_pdf_format(pdf_data, violation_rules_lookup):
    """
    Convert data batch format to PDF generator format.
    Handles multiple images by creating separate PDF data for each image.

    Args:
        pdf_data (dict): Single violation data from the batch
        violation_rules_lookup (dict): Dictionary mapping violation types to regulation info

    Returns:
        list: List of PDF-ready data dictionaries (one per image)
    """
    pdf_ready_list = []

    # # Parse mailing address
    # mail_city, mail_state, mail_zip = parse_city_state_zip(
    #     pdf_data.get("mailing_city_st_zip")
    # )

    # Hard-coded district information (as requested)
    district_info = {
        "district_address_line1": "c/o Public Alliance LLC",
        "district_address_line2": "7555 E. Hampden Ave., Suite 501",
        "district_phone": "(720) 213-6621",
    }

    # Get regulation information
    violation_type = pdf_data.get("violation_type")
    regulation = violation_rules_lookup.get(
        violation_type,
        {
            "code": "N/A",
            "title": violation_type or "Unknown Violation",
            "text": "Regulation text not available for this violation type.",
        },
    )

    # Handle multiple images - create separate PDF data for each image
    violation_images = pdf_data.get("violation_images", [])

    if not violation_images:
        # If no images, create one PDF entry without image
        violation_images = [{"file_path": None}]

    for image_info in violation_images:
        pdf_ready_data = {
            # Required district information
            "district_name": pdf_data.get("district_name", ""),
            **district_info,
            # Date information
            "notice_date": pdf_data.get("violation_date", ""),
            # Homeowner information
            "homeowner_name": pdf_data.get("homeowner_name", ""),
            "homeowner_address_line1": pdf_data.get("mailing_address", ""),
            "homeowner_city": mail_city or "",
            "homeowner_state": mail_state or "",
            "homeowner_zip": mail_zip or "",
            "homeowner_email": pdf_data.get("homeowner_email", ""),
            # Property information
            "property_address": pdf_data.get("property_address", ""),
            # Violation information
            "violation_type": violation_type or "",
            "violation_image_path": image_info.get("file_path", ""),
            # Regulation information
            "regulation": regulation,
            # # Optional fields
            # "homeowner_salutation": extract_salutation(
            #     pdf_data.get("homeowner_name", "")
            # ),
        }

        pdf_ready_list.append(pdf_ready_data)

    return pdf_ready_list


# highland_mead_data = {
#     "district_name": "Highlands Mead Metro District",
# }

# highland_mead_rules = {
#     "1.8 Vehicle Repair": {
#         "code": "1.8",
#         "title": "Vehicle Repair",
#         "text": (
#             "No activity such as, but not limited to, maintenance, repair, rebuilding, dismantling, repainting or "
#             "servicing of any kind of vehicles, trailers or boats, may be performed or conducted in the Community "
#             "unless it is done within a completely enclosed structure that screens the sight and sound of the activity "
#             "from the street, alley, and from adjoining property. The foregoing restriction shall not be deemed to "
#             "prevent the washing and polishing of any motor vehicle, boat, trailer, motor cycle, or other vehicle, "
#             "together with those activities normally incident and necessary to such washing and polishing on a Lot."
#         ),
#     },
#     "2.49 Trash Containers": {
#         "code": "2.49",
#         "title": "Trash Containers",
#         "text": (
#             "Trash containers shall only be placed at curbside for pickup after 6:00 p.m. on the day before pick-up and "
#             "shall be returned to a proper storage location by 9:00 p.m. the day of pick-up. Trash containers shall be "
#             "stored out of sight at all times except on the day of pickup, and shall be kept in a clean and sanitary condition."
#         ),
#     },
#     "2.26 Landscaping": {
#         "code": "2.26",
#         "title": "Landscaping",
#         "text": (
#             "Landscaping must be kept at all times in a neat, healthy, weed-free, and attractive condition."
#         ),
#     },
#     "2.51 Unsightly Conditions": {
#         "code": "2.51",
#         "title": "Unsightly Conditions",
#         "text": (
#             "No unsightly articles or conditions shall be permitted to remain or accumulate on any Lot. By way of example, "
#             "but not limitation, such items could include rock or mulch piles, construction materials, abandoned toys, "
#             "inoperable vehicles, dead or dying landscaping, peeling or faded paint, gardening equipment not in actual use, "
#             "fencing in disrepair, etc."
#         ),
#     },
# }

# # Define field renaming
# rename_map = {
#     "Account Name": "homeowner_name",
#     "MailAddressLine1": "homeowner_address_line1",
#     "MailCity": "homeowner_city",
#     "MailState": "homeowner_state",
#     "MailZip": "homeowner_zip",
#     "Violation Date": "notice_date",
#     "Violation Type": "violation_type",
#     "Email": "homeowner_email",
#     "PropertyAddressLine1": "property_address",
# }


# def read_dataset(
#     filename="../datasets/HighlandsMead5-20Data.xlsx",
# ):
#     """
#     Reads the HighlandsMead Excel file and returns a pandas DataFrame.
#     :param filename: Path to the Excel file (default: datasets/HighlandsMead.xlsx)
#     :return: pandas DataFrame or None if file not found/error
#     """
#     if not os.path.exists(filename):
#         print(f"File not found: {filename}")
#         return None
#     try:
#         df = pd.read_excel(filename)
#         print(f"Loaded {len(df)} rows from {filename}")
#         return df
#     except Exception as e:
#         print(f"Error reading {filename}: {e}")
#         return None


# def clean_address(address):
#     """
#     Removes trailing periods from common street abbreviations (e.g., 'St.' -> 'St').
#     """
#     # List of common abbreviations (add more as needed)
#     abbreviations = [
#         "st",
#         "ave",
#         "blvd",
#         "dr",
#         "rd",
#         "ln",
#         "ct",
#         "pl",
#         "trl",
#         "pkwy",
#         "cir",
#         "ter",
#         "way",
#     ]
#     # Regex: replace abbreviation with period at end of word with no period
#     pattern = re.compile(
#         r"\b(" + "|".join(abbreviations) + r")\.(?=\s|$)", re.IGNORECASE
#     )
#     return pattern.sub(lambda m: m.group(1), address)


# def collect_image(address):
#     """
#     Collects the image path for a given address.
#     :param address: Address string to search for images
#     """

#     # remove all whitespace from address, remove punctuation, convert to lowercase
#     address = clean_address(address)
#     address = "".join(address.split())
#     address = address.lower()

#     # define the directory where images are stored
#     image_dir = "../images/HighlandsMead"

#     # find all images in the directory that match the address
#     images = [
#         os.path.join(image_dir, f)
#         for f in os.listdir(image_dir)
#         if address.lower() in f.lower()
#         and f.lower().endswith((".jpg", ".jpeg", ".png"))
#     ]

#     # return the image page
#     if images:
#         image_path = images[0]  # Take the first matching image
#         print(f"Image found for {address}: {image_path}")
#         return image_path
#     else:
#         print(f"No image found for {address}")
#         return None


# def generate_pdfs():
#     """
#     Function to generate PDFs based on the Highland Mead rules and dataset.
#     This function reads the dataset, applies the rules, and generates PDF files for each address.
#     """

#     df = read_dataset()
#     if df is None:
#         print("No data to process.")
#         return

#     # # collect addrresses for Avery 48807 labels
#     # addresses = []
#     # for index, row in df.iterrows():
#     #     {
#     #         "name": row.get("Account Name", "N/A"),
#     #         "street": row.get("MailAddressLine1", "N/A"),
#     #         "city": row.get("MailCity", "N/A"),
#     #         "state": row.get("MailState", "N/A"),
#     #         "zip": row.get("MailZip", "N/A"),
#     #     }

#     # # send to Avery 48807 labels
#     # create_avery_48807_labels(addresses)

#     # strip trailing whitespace from column names
#     df.columns = df.columns.str.strip()

#     for index, row in df.iterrows():
#         # Filter condition (e.g., only if 'Violation' is not empty)
#         if pd.isna(row.get("Violation Type")) or pd.isna(row.get("Account Name")):
#             continue  # Skip invalid/incomplete rows

#         # Build renamed dictionary
#         data_dict = {
#             new_key: row[old_key]
#             for old_key, new_key in rename_map.items()
#             if pd.notna(row.get(old_key))
#         }

#         # add highland_mead_data to data_dict
#         data_dict.update(highland_mead_data)

#         # match violation type with regulations
#         data_dict["regulation"] = highland_mead_rules[row.get("Code")]

#         # collect image path for violation
#         data_dict["violation_image_path"] = collect_image(data_dict["property_address"])

#         # If no image is found, skip PDF generation for this row
#         if not data_dict["violation_image_path"]:
#             print(
#                 f"Skipping PDF generation for {data_dict['property_address']} (no image found)."
#             )
#             continue

#         # convert notice date to string
#         if pd.notna(data_dict.get("notice_date")):
#             data_dict["notice_date"] = data_dict["notice_date"].strftime("%Y-%m-%d")
#         else:
#             data_dict["notice_date"] = "N/A"

#         # Create PDF generator
#         generator = ViolationNoticePDF()

#         # Generate PDF
#         pdf_path = generator.generate_pdf(data_dict)
#         print(f"PDF generated at: {pdf_path}")


# def create_avery_48807_labels(addresses, output_file="avery_48807_labels.docx"):
#     """
#     Function to create Avery 48807 labels in a Word document.
#     :param addresses: List of address dictionaries with keys: name, street, city, state, zip
#     """
#     doc = Document()

#     labels_per_page = 10  # 2 cols × 5 rows
#     per_row = 2
#     per_col = 5

#     chunks = [
#         addresses[i : i + labels_per_page]
#         for i in range(0, len(addresses), labels_per_page)
#     ]

#     for page in chunks:
#         table = doc.add_table(rows=per_col, cols=per_row)
#         table.alignment = WD_TABLE_ALIGNMENT.CENTER
#         table.autofit = False

#         # Set column widths and height approximations
#         for row in table.rows:
#             for cell in row.cells:
#                 cell.width = Inches(4)
#                 cell.height = Inches(2)
#                 cell.paragraphs[0].paragraph_format.alignment = (
#                     WD_PARAGRAPH_ALIGNMENT.LEFT
#                 )

#         i = 0
#         for row in table.rows:
#             for cell in row.cells:
#                 if i < len(page):
#                     a = page[i]
#                     label = f"{a['name']}\n{a['street']}\n{a['city']}, {a['state']} {a['zip']}"
#                     p = cell.paragraphs[0]
#                     run = p.add_run(label)
#                     run.font.size = Pt(12)
#                     i += 1
#                 else:
#                     cell.text = ""

#         doc.add_page_break()

#     doc.save(output_file)
#     print(f"Saved Avery 48807 label file as: {output_file}")
